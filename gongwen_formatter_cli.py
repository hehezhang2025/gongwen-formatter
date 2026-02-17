#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
公文格式调整工具 - 命令行版本
完全不依赖GUI，纯命令行操作
按照GB/T 9704-2012《党政机关公文格式》标准调整Word文档格式
"""

import os
import sys
import re
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn

# 数字到中文的映射
NUM_TO_CHINESE = {
    1: '一', 2: '二', 3: '三', 4: '四', 5: '五',
    6: '六', 7: '七', 8: '八', 9: '九', 10: '十',
    11: '十一', 12: '十二', 13: '十三', 14: '十四', 15: '十五',
    16: '十六', 17: '十七', 18: '十八', 19: '十九', 20: '二十'
}

# GB/T 9704-2012 标准格式规范
FORMAT_SPECS = {
    'title': {  # 主标题
        'font_name': '方正小标宋简体',
        'font_size': Pt(22),  # 2号字
        'bold': False,
        'alignment': WD_ALIGN_PARAGRAPH.CENTER,
        'space_before': Pt(0),
        'space_after': Pt(0),
        'line_spacing': Pt(35)
    },
    'recipient': {  # 主送机关
        'font_name': '仿宋_GB2312',
        'font_size': Pt(16),  # 3号字
        'bold': False,
        'alignment': WD_ALIGN_PARAGRAPH.LEFT,
        'first_line_indent': Pt(0),  # 顶格
        'line_spacing': Pt(30)
    },
    'heading1': {  # 一级标题：一、
        'font_name': '黑体',
        'font_size': Pt(16),  # 3号字
        'bold': False,
        'first_line_indent': Pt(32),
        'line_spacing': Pt(30),
        'alignment': WD_ALIGN_PARAGRAPH.LEFT  # 左对齐
    },
    'heading2': {  # 二级标题：（一）
        'font_name': '楷体_GB2312',
        'font_size': Pt(16),
        'bold': True,
        'first_line_indent': Pt(32),
        'line_spacing': Pt(30),
        'alignment': WD_ALIGN_PARAGRAPH.LEFT  # 左对齐
    },
    'heading3': {  # 三级标题：1.
        'font_name': '仿宋_GB2312',
        'font_size': Pt(16),
        'bold': True,
        'first_line_indent': Pt(32),
        'line_spacing': Pt(30),
        'alignment': WD_ALIGN_PARAGRAPH.LEFT  # 左对齐
    },
    'heading4': {  # 四级标题：(1)
        'font_name': '仿宋_GB2312',
        'font_size': Pt(16),
        'bold': False,
        'first_line_indent': Pt(32),
        'line_spacing': Pt(30),
        'alignment': WD_ALIGN_PARAGRAPH.LEFT  # 左对齐
    },
    'body': {  # 正文
        'font_name': '仿宋_GB2312',
        'font_size': Pt(16),
        'bold': False,
        'first_line_indent': Pt(32),
        'line_spacing': Pt(30),
        'alignment': WD_ALIGN_PARAGRAPH.LEFT
    },
    'signature': {  # 发文机关署名
        'font_name': '仿宋_GB2312',
        'font_size': Pt(16),
        'bold': False,
        'alignment': WD_ALIGN_PARAGRAPH.RIGHT,
        'right_indent': Pt(32),
        'line_spacing': Pt(30)
    },
    'date': {  # 成文日期
        'font_name': '仿宋_GB2312',
        'font_size': Pt(16),
        'bold': False,
        'alignment': WD_ALIGN_PARAGRAPH.RIGHT,
        'right_indent': Pt(64),
        'line_spacing': Pt(30)
    },
    'caption': {  # 表图说明
        'font_name': '仿宋_GB2312',
        'font_size': Pt(12),  # 5号字，比正文小
        'bold': False,
        'alignment': WD_ALIGN_PARAGRAPH.CENTER,  # 居中
        'line_spacing': Pt(20)
    }
}

def is_title(paragraph, paragraph_count):
    """判断是否是主标题"""
    text = paragraph.text.strip()
    if not text:
        return False
    
    # 生成所有可能的序号前缀
    heading_prefixes = []
    heading_prefixes.extend([f'{NUM_TO_CHINESE[i]}、' for i in range(1, 21)])
    heading_prefixes.extend([f'（{NUM_TO_CHINESE[i]}）' for i in range(1, 21)])
    heading_prefixes.extend([f'{i}.' for i in range(1, 21)])
    heading_prefixes.extend([f'{NUM_TO_CHINESE[i]}是' for i in range(1, 21)])
    
    if any(text.startswith(prefix) for prefix in heading_prefixes):
        return False
    
    # 排除主送机关（以全角冒号结尾）
    if text.endswith('：'):
        return False
    
    # ⭐排除正文开头常用词
    body_start_keywords = ['为', '根据', '按照', '依据', '经', '现', '特']
    if any(text.startswith(kw) for kw in body_start_keywords):
        return False
    
    # ⭐排除附件标记
    if text.startswith('附件') and ('：' in text or '：' in text):
        return False
    
    # 标题通常是前几段，且包含关键词
    if paragraph_count <= 3:
        # 标题通常包含文种词
        title_keywords = ['通知', '报告', '决定', '意见', '办法', '方案', '规定', '通报', '请示', '批复', '函', '纪要']
        if any(kw in text for kw in title_keywords):
            return True
    
    return False

def is_recipient(text):
    """判断是否是主送机关"""
    if not text:
        return False
    # 主送机关：以全角冒号结尾
    if text.endswith('：'):
        # 排除附件标记（附件：开头）
        if text.startswith('附件') and len(text) > 3:
            return False
        # 包含机关关键词或"各"字
        keywords = ['局', '委', '厅', '部', '省', '市', '区', '县', '办', '中心', '公司', '管理', '各']
        if any(kw in text for kw in keywords):
            return True
    return False

def get_heading_level(text):
    """判断段落的标题级别"""
    if not text:
        return None
    
    # 一级标题：一、二、三、（必须是顿号，不是其他标点）
    level1_prefixes = [f'{NUM_TO_CHINESE[i]}、' for i in range(1, 21)]
    if any(text.startswith(prefix) for prefix in level1_prefixes):
        return 1
    
    # 二级标题：（一）（二）（必须是括号，后面不能有标点）
    # ⭐也支持"（一）、"等错误格式（会在apply_paragraph_format中修正）
    level2_prefixes = [f'（{NUM_TO_CHINESE[i]}）' for i in range(1, 21)]
    for prefix in level2_prefixes:
        if text.startswith(prefix):
            return 2
    # 检查错误格式：（一）、、（一）。等
    for i in range(1, 21):
        if text.startswith(f'（{NUM_TO_CHINESE[i]}）、') or text.startswith(f'（{NUM_TO_CHINESE[i]}）。'):
            return 2
    
    # 三级标题：1. 2. 3.（必须是半角点，不是顿号）
    # ⭐也支持"1、"格式（错误格式，会在apply_paragraph_format中修正为"1."）
    for i in range(1, 21):
        if text.startswith(f'{i}.') or text.startswith(f'{i}、'):
            return 3
    
    # 四级标题：(1) (2) (3)（半角括号）
    # ⭐也支持"(1)、" "(1)." "(1)。"等错误格式（会在apply_paragraph_format中修正）
    for i in range(1, 21):
        if text.startswith(f'({i})') or text.startswith(f'({i})、') or text.startswith(f'({i}).') or text.startswith(f'({i})。'):
            return 4
    
    return None

def detect_heading_after_numbering_removed(text):
    """检测移除自动编号后可能的标题（用于举一反三）
    移除自动编号后，原本的一级标题可能变成普通文字，需要通过内容推断
    """
    if not text:
        return None
    
    # 如果已经是标准格式的标题，直接返回
    level = get_heading_level(text)
    if level:
        return level
    
    # ⭐排除正文特征：以常见正文起始词开头
    body_start = ['为', '根据', '按照', '依据', '经', '现', '特', '鉴于', '考虑']
    if any(text.startswith(w) for w in body_start):
        return None
    
    # ⭐排除附件标记和附件列表项
    # 附件：、附件：1.、      2.（6空格开头）等
    import re
    if re.match(r'^附件\d*[：:.]', text):
        return None
    # 排除6个空格开头+数字+点的附件列表项
    if re.match(r'^\s{6}\d+\.', text):
        return None
    
    # ⭐排除以冒号结尾的（正文说明性文字）
    if text.endswith('：') or text.endswith(':'):
        return None
    
    # ⭐排除表格和图片说明（以"表"或"图"开头且包含序号和冒号）
    import re
    if re.match(r'^[表图]\d+[：:]', text):
        return None
    
    # 检查是否像一级标题的特征：
    # 1. 字数适中（通常6-20字，标题不会太长）
    # 2. 不以句号结尾
    # 3. 包含关键动词或主题词
    if 6 <= len(text) <= 20 and not text.endswith('。'):
        # 一级标题常见关键词
        h1_keywords = ['推进', '加强', '提升', '优化', '深化', '强化', '完善', '创新', 
                       '建设', '落实', '实施', '开展', '坚持', '注重', '突出', '聚焦',
                       '治理', '管理', '服务', '保障', '发展', '改革']
        
        if any(kw in text for kw in h1_keywords):
            return 1
    
    return None

def is_attachment_marker(text):
    """判断是否是附件标记"""
    if not text:
        return False
    
    # 附件标记特征：
    # 1. 包含"附件"关键词
    # 2. 可能带序号：附件1、附件一、附件：、附件 1：等
    # 3. 通常是单独一行，不会有其他内容
    
    attachment_patterns = [
        r'^附件[：:\s]*$',  # 单独的"附件"或"附件："
        r'^附件\d+[：:\s]*$',  # 附件1、附件2（后面没有其他内容）
        r'^附件[一二三四五六七八九十]+[：:\s]*$',  # 附件一、附件二（后面没有其他内容）
    ]
    
    for pattern in attachment_patterns:
        if re.search(pattern, text):
            return True
    
    return False

def is_table_or_figure_caption(text):
    """判断是否是表格或图片说明
    格式：表1：XX、表2：XX、图1：XX、图片1：XX等
    """
    if not text:
        return False
    
    import re
    # 匹配：表/图/表格/图片 + 数字 + 冒号
    patterns = [
        r'^表\d+[：:]',      # 表1：、表2：
        r'^图\d+[：:]',      # 图1：、图2：
        r'^表格\d+[：:]',    # 表格1：
        r'^图片\d+[：:]',    # 图片1：
    ]
    
    for pattern in patterns:
        if re.match(pattern, text):
            return True
    
    return False

def is_signature_or_date(paragraphs_list, current_index):
    """判断是否是发文机关署名或成文日期（增强版）"""
    total = len(paragraphs_list)
    
    # 扩大检测范围：最后10个段落都可能是署名/日期
    if current_index < total - 10:
        return None
    
    text = paragraphs_list[current_index].text.strip()
    if not text:
        return None
    
    # 判断日期格式（更宽松，支持XX占位符）
    date_patterns = [
        r'\d{4}年\d{1,2}月\d{1,2}日',
        r'\d{4}年\d{1,2}月XX日',  # 支持XX占位符
        r'[二〇○零一二三四五六七八九十]{4,6}年[一二三四五六七八九十]+月[一二三四五六七八九十]+日',
        r'[二〇○零一二三四五六七八九十]{4,6}年[一二三四五六七八九十]+月XX日',  # 支持XX占位符
    ]
    
    for pattern in date_patterns:
        if re.search(pattern, text):
            return 'date'
    
    # 判断署名（更精确）
    signature_keywords = ['公司', '单位', '部门', '局', '委', '厅', '省', '市', '区', '县', 
                         '中心', '办', '集团', '有限', '科技', '技术', '企业']
    
    # 方法1：当前行包含单位名称，且下一行是日期
    if any(kw in text for kw in signature_keywords):
        # 查找后续几行中是否有日期
        for j in range(1, min(3, total - current_index)):  # 检查后续2行
            next_text = paragraphs_list[current_index + j].text.strip()
            if next_text:  # 跳过空行
                for pattern in date_patterns:
                    if re.search(pattern, next_text):
                        return 'signature'
                break  # 只检查第一个非空行
    
    # 方法2：明确是倒数第二个有效段落（最后一个是日期）
    if current_index == total - 2:
        if any(kw in text for kw in signature_keywords):
            last_text = paragraphs_list[-1].text.strip()
            for pattern in date_patterns:
                if re.search(pattern, last_text):
                    return 'signature'
    
    # 方法3：倒数第三个段落，且倒数第二段是空行，最后一段是日期
    if current_index == total - 3:
        if any(kw in text for kw in signature_keywords):
            # 检查倒数第二段是否为空（已在收集时被过滤）
            last_text = paragraphs_list[-1].text.strip()
            for pattern in date_patterns:
                if re.search(pattern, last_text):
                    return 'signature'
    
    return None

def apply_paragraph_format(paragraph, style_name):
    """应用段落格式"""
    style = FORMAT_SPECS[style_name]
    para_format = paragraph.paragraph_format
    
    # 对齐方式
    if 'alignment' in style:
        para_format.alignment = style['alignment']
    
    # 首行缩进
    if 'first_line_indent' in style:
        para_format.first_line_indent = style['first_line_indent']
    
    # 右缩进
    if 'right_indent' in style:
        para_format.right_indent = style['right_indent']
    
    # 段前段后间距
    if 'space_before' in style:
        para_format.space_before = style['space_before']
    if 'space_after' in style:
        para_format.space_after = style['space_after']
    
    # 行距
    if 'line_spacing' in style:
        para_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        para_format.line_spacing = style['line_spacing']
    
    # ⭐⭐⭐ 清除段落开头的所有空格和Tab（彻底删除）
    # 1. 删除段落开头所有只包含空格/Tab的runs
    # 2. 清理第一个有效run开头的空格/Tab
    while len(paragraph.runs) > 0:
        first_run = paragraph.runs[0]
        # 如果第一个run只包含空格/Tab，直接删除这个run
        if first_run.text and first_run.text.strip() == '':
            # 删除这个只有空格的run
            first_run._element.getparent().remove(first_run._element)
        else:
            # 第一个run有实际内容，清理开头的空格/Tab
            if first_run.text:
                # 同时清理所有空格、Tab和全角空格
                first_run.text = first_run.text.lstrip(' \t\u3000')
            break
    
    # ⭐⭐ 清除所有runs中的Tab字符（标题中可能有自动编号的Tab）
    for run in paragraph.runs:
        if run.text and '\t' in run.text:
            run.text = run.text.replace('\t', '')
    
    # ⭐⭐⭐ 清除标题序号后面的所有空格 + 错误标点 + 多重编号
    # 公文格式规范：
    # - 一级标题：一、（顿号后不加空格，不加其他标点）
    # - 二级标题：（一）（括号后不加空格，不加任何标点，尤其不能加顿号）
    # - 三级标题：1.（半角点后不加空格，不能用顿号）
    # - 四级标题：(1)（半角括号后不加空格，不加任何标点）
    
    # 关键：将所有runs合并处理，然后重新分配（因为Word可能将序号和内容分成不同runs）
    if len(paragraph.runs) > 0:
        import re
        
        # 1. 合并所有runs的文本
        full_text = ''.join([run.text for run in paragraph.runs if run.text])
        original_full_text = full_text
        
        # ⭐⭐⭐ 0. 先清理多重编号（如"（一）1、"、"（二）."等）
        # 使用循环清理机制，最多5次
        for _ in range(5):
            temp = full_text
            
            # 清理二级标题后的三级编号："（一）1、XX" → "（一）XX"、"（一）1.XX" → "（一）XX"
            for i in range(1, 21):
                chinese = ['一','二','三','四','五','六','七','八','九','十','十一','十二','十三','十四','十五','十六','十七','十八','十九','二十'][i-1]
                full_text = re.sub(f'^（{chinese}）\\d+[、.]', f'（{chinese}）', full_text)
            
            # 清理二级标题后的单独点号："（二）.XX" → "（二）XX"
            for i in range(1, 21):
                chinese = ['一','二','三','四','五','六','七','八','九','十','十一','十二','十三','十四','十五','十六','十七','十八','十九','二十'][i-1]
                full_text = re.sub(f'^（{chinese}）\\.', f'（{chinese}）', full_text)
            
            if full_text == temp:
                break
        
        # 2. 对合并后的文本进行清理
        # 一级标题：清除"、"后的空格和错误标点
        for i in range(1, 21):
            # "一、  " → "一、"（清除空格）
            pattern = f'^{NUM_TO_CHINESE[i]}、\\s+'
            replacement = f'{NUM_TO_CHINESE[i]}、'
            full_text = re.sub(pattern, replacement, full_text)
        
        # 二级标题：清除"）"后的空格和错误标点（尤其是顿号）
        for i in range(1, 21):
            # "（一）、" → "（一）"（清除顿号）
            pattern = f'^（{NUM_TO_CHINESE[i]}）、'
            replacement = f'（{NUM_TO_CHINESE[i]}）'
            full_text = re.sub(pattern, replacement, full_text)
            
            # "（一）." → "（一）"（清除半角点）⭐⭐⭐ 新增
            pattern = f'^（{NUM_TO_CHINESE[i]}）\\.'
            replacement = f'（{NUM_TO_CHINESE[i]}）'
            full_text = re.sub(pattern, replacement, full_text)
            
            # "（一）  " → "（一）"（清除空格）
            pattern = f'^（{NUM_TO_CHINESE[i]}）\\s+'
            replacement = f'（{NUM_TO_CHINESE[i]}）'
            full_text = re.sub(pattern, replacement, full_text)
            
            # "（一）。" → "（一）"（清除句号）
            pattern = f'^（{NUM_TO_CHINESE[i]}）。'
            replacement = f'（{NUM_TO_CHINESE[i]}）'
            full_text = re.sub(pattern, replacement, full_text)
        
        # 三级标题：清除点号后的空格，修正错误的顿号
        for i in range(1, 21):
            # "1、" → "1."（修正顿号为半角点）
            pattern = f'^{i}、'
            replacement = f'{i}.'
            full_text = re.sub(pattern, replacement, full_text)
            
            # "1.  " → "1."（清除空格）
            pattern = f'^{i}\\.\\s+'
            replacement = f'{i}.'
            full_text = re.sub(pattern, replacement, full_text)
        
        # 四级标题：清除括号后的空格和错误标点
        for i in range(1, 21):
            # "(1)、" → "(1)"（清除顿号）
            pattern = f'^\\({i}\\)、'
            replacement = f'({i})'
            full_text = re.sub(pattern, replacement, full_text)
            
            # "(1)." → "(1)"（清除点号）
            pattern = f'^\\({i}\\)\\.'
            replacement = f'({i})'
            full_text = re.sub(pattern, replacement, full_text)
            
            # "(1)  " → "(1)"（清除空格）
            pattern = f'^\\({i}\\)\\s+'
            replacement = f'({i})'
            full_text = re.sub(pattern, replacement, full_text)
            
            # "(1)。" → "(1)"（清除句号）
            pattern = f'^\\({i}\\)。'
            replacement = f'({i})'
            full_text = re.sub(pattern, replacement, full_text)
        
        # 3. 如果文本有变化，清空所有runs并用清理后的文本替换
        if full_text != original_full_text:
            # 保留第一个run的格式，删除其他runs
            while len(paragraph.runs) > 1:
                paragraph.runs[-1]._element.getparent().remove(paragraph.runs[-1]._element)
            
            # 将清理后的文本放入第一个run
            if len(paragraph.runs) > 0:
                paragraph.runs[0].text = full_text
    
    # ⭐⭐⭐ 清除标题末尾的标点符号（针对标题样式和title）
    if style_name in ['heading1', 'heading2', 'heading3', 'heading4', 'title']:
        # 需要清除的标点
        punctuation_to_remove = ['。', '；', '，', '.', ';', ',', '、']
        
        # 合并所有runs处理
        if len(paragraph.runs) > 0:
            full_text = ''.join([run.text for run in paragraph.runs if run.text])
            original_full_text = full_text
            
            # 清除末尾标点
            for p in punctuation_to_remove:
                if full_text.rstrip().endswith(p):
                    full_text = full_text.rstrip()[:-1]
                    break
            
            # 如果有变化，更新文本
            if full_text != original_full_text:
                # 保留第一个run的格式，删除其他runs
                while len(paragraph.runs) > 1:
                    paragraph.runs[-1]._element.getparent().remove(paragraph.runs[-1]._element)
                
                # 将清理后的文本放入第一个run
                if len(paragraph.runs) > 0:
                    paragraph.runs[0].text = full_text
    
    # 检查"一是"、"二是"等 - 段落中任意位置
    text = paragraph.text
    shi_prefixes = [f'{NUM_TO_CHINESE[i]}是' for i in range(1, 21)]
    has_shi = any(prefix in text for prefix in shi_prefixes)
    
    # 设置字体格式
    if has_shi:
        # 特殊处理：段落中包含"一是"、"二是"，只加粗这些词（不管是什么类型段落）
        process_shi_paragraph(paragraph, style)
    else:
        # 常规格式
        for run in paragraph.runs:
            run.font.name = style['font_name']
            run._element.rPr.rFonts.set(qn('w:eastAsia'), style['font_name'])
            run.font.size = style['font_size']
            run.font.bold = style['bold']
            run.font.color.rgb = RGBColor(0, 0, 0)
            # ⭐⭐⭐ 清除斜体
            run.font.italic = False

def has_table(paragraph):
    """判断段落是否在表格中"""
    try:
        # 检查段落的父元素是否是表格单元格
        parent = paragraph._element.getparent()
        while parent is not None:
            if parent.tag.endswith('tc'):  # tc = table cell
                return True
            parent = parent.getparent()
        return False
    except:
        return False

def has_image(paragraph):
    """判断段落是否包含图片"""
    try:
        # 检查段落中是否有图片
        for run in paragraph.runs:
            if 'graphic' in run._element.xml or 'pic:pic' in run._element.xml:
                return True
        # 也检查段落的XML
        if 'graphic' in paragraph._element.xml or 'pic:pic' in paragraph._element.xml:
            return True
        return False
    except:
        return False

def center_image_paragraph(paragraph):
    """将包含图片的段落居中对齐"""
    try:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # 清除首行缩进
        paragraph.paragraph_format.first_line_indent = Pt(0)
        paragraph.paragraph_format.left_indent = Pt(0)
        paragraph.paragraph_format.right_indent = Pt(0)
    except:
        pass

def extract_numbering_info(paragraph):
    """提取段落的自动编号信息（级别和ID）"""
    try:
        pPr = paragraph._element.pPr
        if pPr is None:
            return None, None
        
        numPr = pPr.find(qn('w:numPr'))
        if numPr is None:
            return None, None
        
        # 获取编号级别和ID
        ilvl = numPr.find(qn('w:ilvl'))
        numId = numPr.find(qn('w:numId'))
        
        if ilvl is None or numId is None:
            return None, None
        
        level = int(ilvl.get(qn('w:val')))
        num_id = int(numId.get(qn('w:val')))
        
        return level, num_id
    except:
        return None, None

def infer_numbering_text(paragraphs_list, current_index):
    """智能推断应该补回的编号文本
    
    通过分析上下文段落（包括已处理的和未处理的），推断当前段落应该是第几个编号
    """
    try:
        current_para = paragraphs_list[current_index]
        current_level, current_num_id = extract_numbering_info(current_para)
        
        if current_level is None:
            return None
        
        # 生成所有可能的同级别编号前缀
        level_prefixes = []
        if current_level == 0:  # 一级：一、二、三、
            level_prefixes = [f'{NUM_TO_CHINESE[i]}、' for i in range(1, 21)]
        elif current_level == 1:  # 二级：（一）（二）
            level_prefixes = [f'（{NUM_TO_CHINESE[i]}）' for i in range(1, 21)]
        elif current_level == 2:  # 三级：1. 2. 3.
            level_prefixes = [f'{i}.' for i in range(1, 21)]
        
        # 统计当前段落之前有多少个同级别段落
        # 包括：1) 还有自动编号的段落  2) 已经处理过、文本开头有编号的段落
        same_level_count = 0
        
        for i in range(current_index):
            para = paragraphs_list[i]
            text = para.text.strip()
            
            # 方法1：检查是否还有自动编号（同级别、同ID）
            level, num_id = extract_numbering_info(para)
            if level == current_level and num_id == current_num_id:
                same_level_count += 1
                continue
            
            # 方法2：检查文本开头是否有同级别编号（已处理过的段落）
            for prefix in level_prefixes:
                if text.startswith(prefix):
                    same_level_count += 1
                    break
        
        # 当前段落是第几个（从1开始）
        sequence_number = same_level_count + 1
        
        # 根据级别生成对应格式的编号文本
        if current_level == 0:  # 一级标题：一、二、三、
            if sequence_number in NUM_TO_CHINESE:
                return f'{NUM_TO_CHINESE[sequence_number]}、'
        elif current_level == 1:  # 二级标题：（一）（二）
            if sequence_number in NUM_TO_CHINESE:
                return f'（{NUM_TO_CHINESE[sequence_number]}）'
        elif current_level == 2:  # 三级标题：1. 2. 3.
            return f'{sequence_number}.'
        elif current_level == 3:  # 四级标题：(1) (2) (3)
            return f'({sequence_number})'
        
        return None
        
    except:
        return None

def remove_numbering_smart(paragraphs_list, current_index):
    """智能移除段落的自动编号，并根据上下文推断编号补回"""
    try:
        paragraph = paragraphs_list[current_index]
        
        # 先获取当前编号级别
        current_level, current_num_id = extract_numbering_info(paragraph)
        
        # 推断应该补回的编号
        numbering_text = infer_numbering_text(paragraphs_list, current_index)
        
        # 移除编号格式
        pPr = paragraph._element.pPr
        if pPr is not None:
            numPr = pPr.find(qn('w:numPr'))
            if numPr is not None:
                pPr.remove(numPr)
        
        # 如果推断出了编号文本，且段落开头没有该文本，则补回去
        if numbering_text and current_level is not None:
            current_text = paragraph.text.strip()
            
            # 检查是否已经有编号文本了
            has_numbering = False
            for i in range(1, 21):
                if current_level == 0 and current_text.startswith(f'{NUM_TO_CHINESE[i]}、'):
                    has_numbering = True
                    break
                elif current_level == 1 and current_text.startswith(f'（{NUM_TO_CHINESE[i]}）'):
                    has_numbering = True
                    break
                elif current_level == 2 and current_text.startswith(f'{i}.'):
                    has_numbering = True
                    break
                elif current_level == 3 and current_text.startswith(f'({i})'):
                    has_numbering = True
                    break
            
            if not has_numbering:
                # 在段落开头插入编号文本
                if len(paragraph.runs) > 0:
                    first_run = paragraph.runs[0]
                    first_run.text = numbering_text + first_run.text
                else:
                    # 如果没有run，创建一个新的
                    paragraph.add_run(numbering_text)
                
                return numbering_text
        
        return None
        
    except Exception as e:
        return None

def process_shi_paragraph(paragraph, style):
    """处理"一是"、"二是"等段落，对所有"X是"加粗"""
    text = paragraph.text
    shi_prefixes = [f'{NUM_TO_CHINESE[i]}是' for i in range(1, 21)]
    
    # 查找所有"X是"的位置
    shi_positions = []
    for prefix in shi_prefixes:
        pos = 0
        while True:
            pos = text.find(prefix, pos)
            if pos == -1:
                break
            shi_positions.append((pos, prefix))
            pos += len(prefix)
    
    # 如果没找到任何"X是"，返回False
    if not shi_positions:
        return False
    
    # 按位置排序
    shi_positions.sort(key=lambda x: x[0])
    
    # 清空段落并重建
    paragraph.clear()
    
    last_pos = 0
    for pos, prefix in shi_positions:
        # 添加"X是"前面的普通文本
        if pos > last_pos:
            before_text = text[last_pos:pos]
            before_run = paragraph.add_run(before_text)
            before_run.font.name = style['font_name']
            before_run._element.rPr.rFonts.set(qn('w:eastAsia'), style['font_name'])
            before_run.font.size = style['font_size']
            before_run.font.bold = False
            before_run.font.color.rgb = RGBColor(0, 0, 0)
            before_run.font.italic = False  # ⭐清除斜体
        
        # 添加加粗的"X是"
        bold_run = paragraph.add_run(prefix)
        bold_run.font.name = style['font_name']
        bold_run._element.rPr.rFonts.set(qn('w:eastAsia'), style['font_name'])
        bold_run.font.size = style['font_size']
        bold_run.font.bold = True
        bold_run.font.color.rgb = RGBColor(0, 0, 0)
        bold_run.font.italic = False  # ⭐清除斜体
        
        last_pos = pos + len(prefix)
    
    # 添加最后剩余的文本
    if last_pos < len(text):
        remaining = text[last_pos:]
        remaining_run = paragraph.add_run(remaining)
        remaining_run.font.name = style['font_name']
        remaining_run._element.rPr.rFonts.set(qn('w:eastAsia'), style['font_name'])
        remaining_run.font.size = style['font_size']
        remaining_run.font.bold = False
        remaining_run.font.color.rgb = RGBColor(0, 0, 0)
        remaining_run.font.italic = False  # ⭐清除斜体
    
    return True

def validate_and_fix_heading_structure(paragraphs_list):
    """验证并修正标题层级结构
    
    规则：
    1. 编号连续性：一级标题必须从"一、"开始，连续递增（一二三...），不能跳号
    2. 层级合理性：一级标题下只能是二级标题，不能直接跳到三级或四级
    3. 子编号重置：每个一级标题下的二级标题必须从"（一）"开始，三级标题从"1."开始
    4. 附件独立编号：附件部分的标题编号从"一、"重新开始
    """
    print("  🔍 开始验证标题层级结构...")
    
    # ⭐⭐⭐ 第一步：查找附件标记位置
    attachment_start_index = None
    for i, paragraph in enumerate(paragraphs_list):
        text = paragraph.text.strip()
        if is_attachment_marker(text):
            attachment_start_index = i
            print(f"  📎 检测到附件标记位置: 第{i+1}段，附件内容将独立编号")
            break
    
    # 收集所有标题段落及其级别
    heading_info = []  # [(paragraph, level, current_number, index, is_in_attachment), ...]
    
    for i, paragraph in enumerate(paragraphs_list):
        text = paragraph.text.strip()
        level = get_heading_level(text)
        
        # ⭐⭐⭐ 关键：只有已经有明确编号的才算标题
        # 不使用detect_heading_after_numbering_removed，避免把文档标题误判为一级标题
        if level:
            # 提取当前编号
            current_num = extract_current_number(text, level)
            # 判断是否在附件部分
            is_in_attachment = attachment_start_index is not None and i > attachment_start_index
            heading_info.append((paragraph, level, current_num, i, is_in_attachment))
    
    if not heading_info:
        print("  ✓ 未检测到标题，跳过层级检查")
        return
    
    print(f"  📊 检测到 {len(heading_info)} 个标题")
    
    # ⭐⭐⭐ 第一轮：检测并修正层级跳跃和编号错误
    level_counters = {1: 0, 2: 0, 3: 0, 4: 0}
    last_level = 0
    fixed_count = 0
    has_level2_under_current_level1 = False  # 当前一级标题下是否已经有二级标题
    last_is_in_attachment = False  # 上一个标题是否在附件中
    
    for idx, (paragraph, level, current_num, para_idx, is_in_attachment) in enumerate(heading_info):
        text = paragraph.text.strip()
        original_level = level
        
        # ⭐⭐⭐ 关键：进入附件部分时，重置所有计数器
        if is_in_attachment and not last_is_in_attachment:
            print(f"  📎 进入附件部分，标题编号重新开始")
            level_counters = {1: 0, 2: 0, 3: 0, 4: 0}
            last_level = 0
            has_level2_under_current_level1 = False
        
        # 规则1: 检查层级跳跃
        if last_level > 0 and level > last_level + 1:
            print(f"  ⚠️  层级跳跃: 第{para_idx+1}段 从{last_level}级直接跳到{level}级")
            print(f"      内容: {text[:40]}")
            level = last_level + 1
        
        # 规则2: 检查3级标题是否应该降级
        # 如果当前一级标题下还没有二级标题，3级应该降为2级
        if level == 3 and not has_level2_under_current_level1:
            print(f"  ⚠️  层级错误: 第{para_idx+1}段 三级标题在一级标题下，应降为二级")
            print(f"      内容: {text[:40]}")
            level = 2
        
        # 更新标志
        if level == 1:
            # 新的一级标题，重置标志
            has_level2_under_current_level1 = False
        elif level == 2 and original_level == 2:
            # ⭐只有原本就是二级标题的，才算真正的二级标题
            # 从三级降级来的不算
            has_level2_under_current_level1 = True
        
        # 规则3: 子标题编号重置
        # 当出现同级或更高级别标题时，重置下级计数器
        if level <= last_level:
            for l in range(level + 1, 5):
                level_counters[l] = 0
        
        # 增加当前层级计数
        level_counters[level] += 1
        expected_num = level_counters[level]
        
        # 执行修正
        if original_level != level or current_num != expected_num:
            # 需要修正层级或编号
            if original_level != level:
                print(f"    🔧 第{para_idx+1}段: {original_level}级→{level}级")
            if current_num != expected_num:
                print(f"    🔧 第{para_idx+1}段: 编号{get_number_display(current_num, original_level)}→{get_number_display(expected_num, level)}")
            
            # 执行修正
            fix_heading_number(paragraph, level, expected_num)
            fixed_count += 1
        
        last_level = level
        last_is_in_attachment = is_in_attachment
    
    if fixed_count > 0:
        print(f"  ✅ 共修正 {fixed_count} 个标题\n")
    else:
        print("  ✅ 标题层级结构正确\n")

def normalize_attachment_list(paragraphs_list):
    """规范化附件列表格式
    
    规则：
    1. 第一行：附件：1.XXX
    2. 后续行：      2.XXX（前面6个空格，和1.对齐）
    3. 编号连续性：1、2、3，不能跳号
    4. 格式：仿宋16磅，不加粗（正文格式）
    """
    import re
    
    print("  🔍 开始规范化附件列表...")
    
    # 查找附件列表起始位置
    attachment_list_start = -1
    for i, para in enumerate(paragraphs_list):
        text = para.text.strip()
        # 匹配"附件："或"附件:"开头的行
        if re.match(r'^附件[：:]\s*\d+[、，.]', text):
            attachment_list_start = i
            print(f"  📎 检测到附件列表起始: 第{i+1}段")
            break
    
    if attachment_list_start == -1:
        print("  ✓ 未检测到附件列表\n")
        return
    
    # 收集附件列表项
    attachment_items = []
    current_index = attachment_list_start
    
    # 第一行：提取"附件：1、XXX"中的内容
    first_text = paragraphs_list[current_index].text.strip()
    match = re.match(r'^附件[：:]\s*(\d+)[、，.](.+)$', first_text)
    if match:
        num = int(match.group(1))
        content = match.group(2).strip()
        attachment_items.append((paragraphs_list[current_index], num, content, True))  # True表示是第一行
        current_index += 1
    
    # 后续行：匹配"2、XXX"或"  2、XXX"
    while current_index < len(paragraphs_list):
        text = paragraphs_list[current_index].text.strip()
        
        # 匹配数字+顿号/逗号/点开头
        match = re.match(r'^(\d+)[、，.](.+)$', text)
        if match:
            num = int(match.group(1))
            content = match.group(2).strip()
            attachment_items.append((paragraphs_list[current_index], num, content, False))  # False表示不是第一行
            current_index += 1
        else:
            # 不再是附件列表项，退出
            break
    
    if not attachment_items:
        print("  ✓ 未检测到附件列表项\n")
        return
    
    print(f"  📊 检测到 {len(attachment_items)} 个附件")
    
    # 检查编号连续性并修正
    fixed_count = 0
    for idx, (paragraph, current_num, content, is_first) in enumerate(attachment_items, 1):
        expected_num = idx
        
        # ⭐⭐⭐ 格式规则：
        # 第一行：附件：1.内容
        # 后续行：      2.内容（6个空格，和1.对齐）
        if is_first:
            new_text = f"附件：{expected_num}.{content}"
        else:
            new_text = f"      {expected_num}.{content}"  # 6个空格
        
        if current_num != expected_num:
            print(f"    🔧 第{paragraphs_list.index(paragraph)+1}段: 编号{current_num}→{expected_num}")
            fixed_count += 1
        
        # 更新段落文本
        # 保留第一个run，删除其他runs
        while len(paragraph.runs) > 1:
            paragraph.runs[-1]._element.getparent().remove(paragraph.runs[-1]._element)
        
        if len(paragraph.runs) > 0:
            paragraph.runs[0].text = new_text
        else:
            paragraph.add_run(new_text)
        
        # 清除加粗格式（确保是正文格式）
        for run in paragraph.runs:
            run.font.bold = False
        
        # ⭐⭐⭐ 设置段落格式：左对齐，无缩进
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Pt
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        paragraph.paragraph_format.first_line_indent = Pt(0)
        paragraph.paragraph_format.left_indent = Pt(0)
    
    if fixed_count > 0:
        print(f"  ✅ 共修正 {fixed_count} 个附件编号\n")
    else:
        print("  ✅ 附件列表格式正确\n")


def extract_current_number(text, level):
    """从标题文本中提取当前编号"""
    import re
    
    if level == 1:
        # 一、二、三、
        for i in range(1, 21):
            if text.startswith(f'{NUM_TO_CHINESE[i]}、'):
                return i
    
    elif level == 2:
        # （一）（二）
        for i in range(1, 21):
            if text.startswith(f'（{NUM_TO_CHINESE[i]}）'):
                return i
    
    elif level == 3:
        # 1. 2. 3.
        match = re.match(r'^(\d+)\.', text)
        if match:
            return int(match.group(1))
    
    elif level == 4:
        # (1) (2) (3)
        match = re.match(r'^\((\d+)\)', text)
        if match:
            return int(match.group(1))
    
    return 1  # 默认返回1

def get_number_display(num, level):
    """获取编号的显示文本"""
    if level == 1:
        return f'{NUM_TO_CHINESE.get(num, str(num))}、'
    elif level == 2:
        return f'（{NUM_TO_CHINESE.get(num, str(num))}）'
    elif level == 3:
        return f'{num}.'
    elif level == 4:
        return f'({num})'
    return str(num)

def fix_heading_number(paragraph, level, correct_number):
    """修正标题编号（支持跨层级转换）"""
    import re
    
    if len(paragraph.runs) == 0:
        return
    
    # 合并所有runs的文本
    full_text = ''.join([run.text for run in paragraph.runs if run.text])
    
    # ⭐⭐⭐ 关键改进：循环移除所有编号格式，直到没有任何编号为止
    # 这样可以处理"（一）1、"这种多重编号的情况
    max_iterations = 5  # 最多循环5次，避免死循环
    for _ in range(max_iterations):
        original = full_text
        
        # 移除一级：X、（中文数字+顿号）
        full_text = re.sub(r'^[一二三四五六七八九十]{1,2}、\s*', '', full_text)
        
        # 移除二级：（X）（括号+中文数字+括号）
        full_text = re.sub(r'^（[一二三四五六七八九十]{1,2}）\s*', '', full_text)
        
        # 移除三级变体1：X.（数字+点）
        full_text = re.sub(r'^\d+\.\s*', '', full_text)
        
        # 移除三级变体2：X、（数字+顿号）⭐⭐⭐ 这个之前漏了！
        full_text = re.sub(r'^\d+、\s*', '', full_text)
        
        # 移除四级变体1：(X)（半角括号+数字+半角括号）
        full_text = re.sub(r'^\(\d+\)\s*', '', full_text)
        
        # 移除四级变体2：(X).（半角括号+数字+半角括号+点）⭐⭐⭐ 这个之前也漏了！
        full_text = re.sub(r'^\(\d+\)\.\s*', '', full_text)
        
        # 移除四级变体3：（X）（全角括号+数字+全角括号）
        full_text = re.sub(r'^（\d+）\s*', '', full_text)
        
        # 移除可能的多余点和空格
        full_text = re.sub(r'^\.\s*', '', full_text)
        full_text = re.sub(r'^．\s*', '', full_text)  # 全角点
        
        # 如果没有变化，说明已经清理干净了
        if full_text == original:
            break
    
    # 移除开头的多余空格
    full_text = full_text.lstrip()
    
    # 根据目标层级添加正确的编号
    if level == 1:
        new_text = f'{NUM_TO_CHINESE.get(correct_number, str(correct_number))}、{full_text}'
    elif level == 2:
        new_text = f'（{NUM_TO_CHINESE.get(correct_number, str(correct_number))}）{full_text}'
    elif level == 3:
        new_text = f'{correct_number}.{full_text}'
    elif level == 4:
        new_text = f'({correct_number}){full_text}'
    else:
        return
    
    # 更新段落文本
    # 保留第一个run，删除其他runs
    while len(paragraph.runs) > 1:
        paragraph.runs[-1]._element.getparent().remove(paragraph.runs[-1]._element)
    
    if len(paragraph.runs) > 0:
        paragraph.runs[0].text = new_text

def format_document(input_path):
    """格式化公文文档（完整版）"""
    try:
        print(f"\n📄 正在处理: {os.path.basename(input_path)}")
        print("━" * 50)
        
        # 1. 打开文档
        print("  ⏳ 读取文档...")
        doc = Document(input_path)
        
        # 2. 设置页边距（GB/T 9704-2012标准）
        section = doc.sections[0]
        section.top_margin = Cm(3.7)
        section.bottom_margin = Cm(3.5)
        section.left_margin = Cm(2.8)
        section.right_margin = Cm(2.6)
        print("  ✅ 页边距: 上3.7cm 下3.5cm 左2.8cm 右2.6cm")
        
        # ⭐ 新增：删除文档末尾的空行
        removed_trailing = 0
        while len(doc.paragraphs) > 0:
            last_para = doc.paragraphs[-1]
            if not last_para.text.strip() and not has_image(last_para):
                # 删除最后一个空段落
                p = last_para._element
                p.getparent().remove(p)
                removed_trailing += 1
            else:
                break
        if removed_trailing > 0:
            print(f"  🧹 删除文档末尾空行: {removed_trailing} 个")
        
        # 统计表格和图片
        table_count = len(doc.tables)
        if table_count > 0:
            print(f"  📊 检测到 {table_count} 个表格（将跳过不处理）")
        
        # 3. 收集所有段落（包括空段落，用于智能推断编号）
        all_paragraphs = list(doc.paragraphs)
        
        # 4. 第一遍：智能推断所有编号（在移除之前）
        print("  🔧 智能推断编号文本...")
        numbering_map = {}  # 存储每个段落应该补回的编号文本
        
        for i in range(len(all_paragraphs)):
            numbering_text = infer_numbering_text(all_paragraphs, i)
            if numbering_text:
                numbering_map[i] = numbering_text
        
        print(f"  ✅ 推断出 {len(numbering_map)} 个编号")
        
        # 5. 第二遍：移除所有自动编号并补回
        print("  🔧 移除自动编号并补回...")
        numbering_补回_count = 0
        
        for i, paragraph in enumerate(all_paragraphs):
            # 检查是否有编号格式
            pPr = paragraph._element.pPr
            has_numbering = False
            if pPr is not None:
                numPr = pPr.find(qn('w:numPr'))
                if numPr is not None:
                    has_numbering = True
                    # 移除编号格式
                    pPr.remove(numPr)
            
            # 如果有编号格式，且推断出了编号文本，则补回
            if has_numbering and i in numbering_map:
                numbering_text = numbering_map[i]
                current_text = paragraph.text.strip()
                
                # 检查是否已经有编号文本
                if not current_text.startswith(numbering_text):
                    # 在段落开头插入编号文本
                    if len(paragraph.runs) > 0:
                        first_run = paragraph.runs[0]
                        first_run.text = numbering_text + first_run.text
                    else:
                        paragraph.add_run(numbering_text)
                    
                    print(f"    🔧 补回编号: {numbering_text}")
                    numbering_补回_count += 1
        
        if numbering_补回_count > 0:
            print(f"  ✅ 共补回 {numbering_补回_count} 个编号")
        
        # 6. 删除标题上方的空行（可能有多个连续空行，需要多次扫描）
        print("  🧹 检查并删除标题上方的空行...")
        removed_empty = 0
        
        # 多次扫描，直到没有可删除的空行
        while True:
            deleted_in_this_round = 0
            i = 1  # 从第二段开始检查
            
            while i < len(doc.paragraphs):
                prev_para = doc.paragraphs[i - 1]
                curr_para = doc.paragraphs[i]
                
                prev_text = prev_para.text.strip()
                curr_text = curr_para.text.strip()
                
                # ⭐检查当前段落是否是标题（包括标准格式和智能推断）
                is_heading = get_heading_level(curr_text) is not None
                if not is_heading:
                    # 也检查智能推断的标题
                    is_heading = detect_heading_after_numbering_removed(curr_text) is not None
                
                # 如果当前是标题，且上一段为空，删除上一段
                if is_heading and not prev_text and not has_image(prev_para):
                    p = prev_para._element
                    p.getparent().remove(p)
                    removed_empty += 1
                    deleted_in_this_round += 1
                    # 删除后重新开始循环
                    break
                
                i += 1
            
            # 如果这一轮没有删除任何空行，退出
            if deleted_in_this_round == 0:
                break
        
        if removed_empty > 0:
            print(f"  ✓ 删除标题上方空行: {removed_empty} 个")
        
        # 5. 收集所有非空段落用于格式化
        paragraphs_list = [p for p in doc.paragraphs if p.text.strip() or has_image(p)]
        total = len(paragraphs_list)
        print(f"  📝 共 {total} 个有效段落")
        
        # ⭐⭐⭐ 新增：验证并修正标题层级结构
        validate_and_fix_heading_structure(paragraphs_list)
        
        # ⭐⭐⭐ 新增：规范化附件列表格式
        normalize_attachment_list(paragraphs_list)
        
        # 5. 检测附件位置
        attachment_start_index = None
        for i, paragraph in enumerate(paragraphs_list):
            text = paragraph.text.strip()
            if is_attachment_marker(text):
                attachment_start_index = i
                print(f"  📎 检测到附件标记: {text}（第{i+1}个段落）")
                break
        
        # 6. 处理每个段落
        title_found = False
        recipient_found = False
        paragraph_count = 0
        skipped_table = 0
        processed_image = 0
        in_attachment = False  # 是否进入附件部分
        attachment_title_found = False  # 附件中是否找到标题
        in_attachment_list = False  # 是否在附件列表中（"附件：1、XX  2、XX"）
        
        for paragraph in paragraphs_list:
            text = paragraph.text.strip()
            paragraph_count += 1
            current_index = paragraphs_list.index(paragraph)
            
            # 检查是否进入附件部分
            if attachment_start_index is not None and current_index >= attachment_start_index:
                if not in_attachment:
                    in_attachment = True
                    print(f"\n  📎 === 开始处理附件部分 ===")
            
            # 规则1：跳过表格中的段落
            if has_table(paragraph):
                skipped_table += 1
                continue
            
            # 规则2：图片段落只居中，不做其他处理
            if has_image(paragraph):
                center_image_paragraph(paragraph)
                processed_image += 1
                print(f"  🖼️  图片: 已居中对齐")
                continue
            
            # 附件部分的处理逻辑
            if in_attachment:
                # 附件标记本身：左上角顶格、3号黑体
                if is_attachment_marker(text):
                    # ⭐附件标记：左对齐顶格
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    paragraph.paragraph_format.first_line_indent = Pt(0)  # 顶格
                    paragraph.paragraph_format.left_indent = Pt(0)
                    
                    for run in paragraph.runs:
                        run.font.name = '黑体'
                        if run._element.rPr is not None:
                            run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
                        run.font.size = Pt(16)  # 3号字
                        run.font.bold = True
                        run.font.italic = False  # 清除斜体
                        run.font.color.rgb = RGBColor(0, 0, 0)
                    
                    # ⭐在附件标记前插入分页符（换页）
                    if current_index > 0:
                        from docx.oxml import OxmlElement
                        from docx.oxml.ns import qn as qn_ns
                        # 在当前段落的第一个run前插入分页符
                        if paragraph.runs:
                            run = paragraph.runs[0]
                        else:
                            run = paragraph.add_run()
                        
                        # 创建分页符元素
                        br = OxmlElement('w:br')
                        br.set(qn('w:type'), 'page')
                        
                        # 插入到run的开头
                        run._element.insert(0, br)
                    
                    print(f"  📎 附件标记（换页）: {text[:30]}")
                    continue
                
                # 附件中的主标题（附件的文档标题）
                # ⭐改进：不使用is_title判断，而是检查是否是附件标记后的第一个非空段落
                if not attachment_title_found:
                    # 检查是否像标题（包含文种词或者字数较短）
                    title_keywords = ['通知', '报告', '决定', '意见', '办法', '方案', '规定', '通报', '请示', '批复', '函', '纪要', '制度', '汇编', '计划', '总结']
                    is_likely_title = any(kw in text for kw in title_keywords) or len(text) <= 30
                    
                    # 排除一级标题格式
                    has_standard_heading = get_heading_level(text) is not None
                    
                    if is_likely_title and not has_standard_heading:
                        apply_paragraph_format(paragraph, 'title')
                        print(f"  📌 [附件]标题: {text[:30]}...")
                        attachment_title_found = True
                        continue
                
                # 附件中的标题级别判断
                heading_level = get_heading_level(text)
                if not heading_level:
                    heading_level = detect_heading_after_numbering_removed(text)
                
                if heading_level == 1:
                    apply_paragraph_format(paragraph, 'heading1')
                    print(f"  🔹 [附件]一级标题: {text[:30]}")
                elif heading_level == 2:
                    apply_paragraph_format(paragraph, 'heading2')
                    print(f"    🔸 [附件]二级标题: {text[:30]}")
                elif heading_level == 3:
                    apply_paragraph_format(paragraph, 'heading3')
                    print(f"      ▪️  [附件]三级标题: {text[:30]}")
                elif heading_level == 4:
                    apply_paragraph_format(paragraph, 'heading4')
                    print(f"        • [附件]四级标题: {text[:30]}")
                else:
                    apply_paragraph_format(paragraph, 'body')
                    if paragraph_count % 10 == 0:
                        print(f"  ✓ [附件]已处理 {paragraph_count}/{total} 个段落")
                
                continue
            
            # 正文部分的处理逻辑（原有逻辑）
            # 判断主标题
            if not title_found and is_title(paragraph, paragraph_count):
                apply_paragraph_format(paragraph, 'title')
                print(f"  📌 标题: {text[:30]}...")
                title_found = True
                
                # ⭐⭐⭐ 标题和主送机关之间需要空一行
                # 检查下一段是否是主送机关
                if current_index + 1 < len(paragraphs_list):
                    next_para = paragraphs_list[current_index + 1]
                    next_text = next_para.text.strip()
                    if is_recipient(next_text):
                        # 在标题后插入一个空行
                        # 获取标题段落在文档中的位置
                        title_element = paragraph._element
                        parent = title_element.getparent()
                        title_idx = list(parent).index(title_element)
                        # 在标题后插入空段落
                        from docx.oxml import OxmlElement
                        new_p = OxmlElement('w:p')
                        parent.insert(title_idx + 1, new_p)
                        print(f"  ✓ 在标题和主送机关之间插入空行")
                
                continue
            
            # 判断主送机关（在标题之后）
            if title_found and not recipient_found and is_recipient(text):
                apply_paragraph_format(paragraph, 'recipient')
                print(f"  📨 主送机关: {text[:30]}")
                recipient_found = True
                continue
            
            # 判断署名和日期（优先级提高，在标题判断之前）
            sig_or_date = is_signature_or_date(paragraphs_list, current_index)
            if sig_or_date == 'signature':
                apply_paragraph_format(paragraph, 'signature')
                print(f"  ✍️  署名: {text[:30]}")
                continue
            elif sig_or_date == 'date':
                apply_paragraph_format(paragraph, 'date')
                print(f"  📅 日期: {text[:30]}")
                continue
            
            # ⭐判断表图说明（在标题判断之前）
            if is_table_or_figure_caption(text):
                apply_paragraph_format(paragraph, 'caption')
                print(f"  📊 表图说明: {text[:30]}")
                continue
            
            # ⭐判断附件列表（已规范化格式）
            import re
            # ⭐⭐⭐ 关键：使用原始文本（不strip），保留前导空格
            raw_text = paragraph.text
            
            # 格式1: "附件：1.XX"（第一行）
            if re.match(r'^附件[：:]\d+\.', text):
                # ⭐不能用apply_paragraph_format，因为它会删除前导空格
                # 直接设置格式
                
                # 段落格式：和正文一样，首行缩进2字符
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                paragraph.paragraph_format.first_line_indent = Pt(32)  # 2字符缩进
                paragraph.paragraph_format.left_indent = Pt(0)
                paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
                paragraph.paragraph_format.line_spacing = Pt(28)
                
                # 字体格式（仿宋16磅）
                for run in paragraph.runs:
                    run.font.name = '仿宋_GB2312'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
                    run.font.size = Pt(16)
                    run.font.bold = False
                    run.font.color.rgb = RGBColor(0, 0, 0)
                
                print(f"  📎 附件列表第一项: {text[:30]}")
                continue
            
            # 格式2: "      2.XX"（后续行，6个空格开头）
            # ⭐使用raw_text检测前导空格
            if re.match(r'^\s{6}\d+\.', raw_text):
                # ⭐不能用apply_paragraph_format，因为它会删除前导空格
                # 直接设置格式
                
                # 段落格式：和正文一样，首行缩进2字符
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                paragraph.paragraph_format.first_line_indent = Pt(32)  # 2字符缩进
                paragraph.paragraph_format.left_indent = Pt(0)
                paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
                paragraph.paragraph_format.line_spacing = Pt(28)
                
                # 字体格式（仿宋16磅）
                for run in paragraph.runs:
                    run.font.name = '仿宋_GB2312'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
                    run.font.size = Pt(16)
                    run.font.bold = False
                    run.font.color.rgb = RGBColor(0, 0, 0)
                
                print(f"  📎 附件列表项: {text.strip()[:30]}")
                continue
            
            # 旧版附件列表检测（向后兼容，已废弃）
            # 检测"附件："开头的段落
            if text.startswith('附件') and ('：' in text or ':' in text):
                in_attachment_list = True
                apply_paragraph_format(paragraph, 'body')
                print(f"  📎 附件列表起始: {text[:30]}")
                continue
            
            # 如果在附件列表中，检测列表项（如"2、XX"、"  2、XX"等）
            if in_attachment_list:
                # 判断是否是列表项：以数字+顿号开头，或前面有空格缩进
                is_list_item = re.match(r'^\s*\d+[、，]', text)
                if is_list_item:
                    apply_paragraph_format(paragraph, 'body')
                    print(f"  📎 附件列表项: {text[:30]}")
                    continue
                else:
                    # 不再是列表项，退出附件列表状态
                    in_attachment_list = False
            
            # 判断标题级别（先用标准格式，再用智能推断）
            heading_level = get_heading_level(text)
            if not heading_level:
                # 移除自动编号后，可能需要智能推断
                heading_level = detect_heading_after_numbering_removed(text)
            
            if heading_level == 1:
                apply_paragraph_format(paragraph, 'heading1')
                print(f"  🔹 一级标题: {text[:30]}")
            elif heading_level == 2:
                apply_paragraph_format(paragraph, 'heading2')
                print(f"    🔸 二级标题: {text[:30]}")
            elif heading_level == 3:
                apply_paragraph_format(paragraph, 'heading3')
                print(f"      ▪️  三级标题: {text[:30]}")
            elif heading_level == 4:
                apply_paragraph_format(paragraph, 'heading4')
                print(f"        • 四级标题: {text[:30]}")
            else:
                apply_paragraph_format(paragraph, 'body')
                if paragraph_count % 10 == 0:
                    print(f"  ✓ 已处理 {paragraph_count}/{total} 个段落")
        
        print(f"  ✓ 全部 {total} 个段落处理完成")
        
        # 统计信息
        if skipped_table > 0:
            print(f"  ⏭️  跳过表格段落: {skipped_table} 个")
        if processed_image > 0:
            print(f"  🖼️  处理图片: {processed_image} 个（已居中）")
        
        # 5. 保存文档
        dir_name = os.path.dirname(input_path)
        base_name = os.path.basename(input_path)
        output_path = os.path.join(dir_name, f"done_{base_name}")
        
        print(f"  💾 保存文档...")
        doc.save(output_path)
        
        print("━" * 50)
        print(f"✅ 处理完成！")
        print(f"📁 输出文件: {output_path}\n")
        
        return True
        
    except Exception as e:
        print(f"❌ 处理失败: {str(e)}\n")
        import traceback
        traceback.print_exc()
        return False

def main():
    """
    主函数 - 命令行交互
    """
    print("\n" + "=" * 50)
    print("  📄 公文格式调整工具（命令行版）")
    print("=" * 50)
    print("\n使用方法：")
    print("  1. 拖拽Word文档到此窗口")
    print("  2. 按回车键开始处理")
    print("  3. 输入 'q' 退出程序")
    print("\n" + "=" * 50 + "\n")
    
    while True:
        try:
            # 获取用户输入
            user_input = input("📎 请拖拽Word文档到此处（或输入q退出）: ").strip()
            
            # 退出
            if user_input.lower() == 'q':
                print("\n👋 再见！\n")
                break
            
            # 处理路径（支持多种格式）
            # 1. 去除首尾的引号（单引号或双引号）
            file_path = user_input.strip('"').strip("'").strip()
            
            # 2. 处理macOS拖拽时的反斜杠转义（如：测试\ 文件.docx）
            # 将 "\ " 替换为 " "（空格前的反斜杠是转义符）
            file_path = file_path.replace('\\ ', ' ')
            
            # 3. 处理其他常见的转义字符
            file_path = file_path.replace('\\(', '(').replace('\\)', ')')
            file_path = file_path.replace('\\[', '[').replace('\\]', ']')
            file_path = file_path.replace('\\&', '&')
            
            # 检查文件
            if not file_path:
                continue
                
            if not os.path.exists(file_path):
                print(f"❌ 文件不存在: {file_path}\n")
                continue
            
            if not file_path.lower().endswith('.docx'):
                print("❌ 只支持.docx格式的文件\n")
                continue
            
            # 处理文档
            success = format_document(file_path)
            
            if success:
                print("━" * 50)
                print("✨ 可以继续处理下一个文档")
                print("━" * 50 + "\n")
            
        except KeyboardInterrupt:
            print("\n\n👋 程序已终止\n")
            break
        except Exception as e:
            print(f"\n❌ 发生错误: {str(e)}\n")

if __name__ == '__main__':
    main()
