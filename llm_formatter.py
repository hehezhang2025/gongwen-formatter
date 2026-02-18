#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
LLM 增强格式化模块
使用本地 Qwen 模型智能识别文档结构，然后应用格式
"""

import os
import sys
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn

from llm_client import OllamaClient
from gongwen_formatter_cli import apply_paragraph_format, has_table, has_image, center_image_paragraph


def validate_llm_result(llm_result, doc):
    """验证 LLM 识别结果的有效性"""
    try:
        # 检查必需字段
        if "paragraphs" not in llm_result:
            print("  ❌ LLM结果缺少 'paragraphs' 字段")
            return False
        
        paragraphs = llm_result["paragraphs"]
        if not isinstance(paragraphs, list):
            print("  ❌ 'paragraphs' 不是列表")
            return False
        
        if len(paragraphs) == 0:
            print("  ❌ 'paragraphs' 为空")
            return False
        
        # 检查每个段落的格式
        valid_types = {
            'title', 'recipient', 'heading1', 'heading2', 'heading3', 'heading4',
            'body', 'attachment_marker', 'signature', 'date'
        }
        
        for i, para in enumerate(paragraphs):
            if not isinstance(para, dict):
                print(f"  ❌ 第 {i} 个段落不是字典")
                return False
            
            if "type" not in para:
                print(f"  ❌ 第 {i} 个段落缺少 'type' 字段")
                return False
            
            if para["type"] not in valid_types:
                print(f"  ⚠️  第 {i} 个段落类型无效: {para['type']}，将视为body")
                para["type"] = "body"  # 自动修正
            
            if "index" not in para:
                print(f"  ❌ 第 {i} 个段落缺少 'index' 字段")
                return False
        
        print("  ✅ LLM结果验证通过")
        return True
        
    except Exception as e:
        print(f"  ❌ LLM结果验证失败: {str(e)}")
        return False


def apply_formats_by_llm(doc, llm_result):
    """根据 LLM 识别结果应用格式（只改格式，不改内容）"""
    
    # 创建段落索引映射（只包含非空段落）
    para_map = {}
    valid_para_count = 0
    
    for para in doc.paragraphs:
        # 跳过表格中的段落
        if has_table(para):
            continue
        
        # 跳过图片段落（但要居中）
        if has_image(para):
            center_image_paragraph(para)
            continue
        
        text = para.text.strip()
        if text:
            para_map[valid_para_count] = para
            valid_para_count += 1
    
    print(f"  📊 文档共有 {valid_para_count} 个有效段落")
    print(f"  🤖 LLM识别了 {len(llm_result['paragraphs'])} 个段落")
    
    # 统计各类型数量
    type_counts = {}
    
    # 遍历 LLM 识别结果
    for item in llm_result['paragraphs']:
        index = item.get('index')
        para_type = item.get('type', 'body')
        
        # 统计
        type_counts[para_type] = type_counts.get(para_type, 0) + 1
        
        # 检查索引是否有效
        if index is None or index not in para_map:
            continue
        
        paragraph = para_map[index]
        
        # 根据类型应用格式
        if para_type == 'title':
            apply_paragraph_format(paragraph, 'title')
            print(f"  📌 标题: {paragraph.text[:30]}...")
        
        elif para_type == 'recipient':
            apply_paragraph_format(paragraph, 'recipient')
            print(f"  📨 主送机关: {paragraph.text[:30]}")
        
        elif para_type == 'heading1':
            apply_paragraph_format(paragraph, 'heading1')
            print(f"  🔹 一级标题: {paragraph.text[:30]}")
        
        elif para_type == 'heading2':
            apply_paragraph_format(paragraph, 'heading2')
            print(f"    🔸 二级标题: {paragraph.text[:30]}")
        
        elif para_type == 'heading3':
            apply_paragraph_format(paragraph, 'heading3')
            print(f"      ▪️  三级标题: {paragraph.text[:30]}")
        
        elif para_type == 'heading4':
            apply_paragraph_format(paragraph, 'heading4')
            print(f"        • 四级标题: {paragraph.text[:30]}")
        
        elif para_type == 'signature':
            apply_paragraph_format(paragraph, 'signature')
            print(f"  ✍️  署名: {paragraph.text[:30]}")
        
        elif para_type == 'date':
            apply_paragraph_format(paragraph, 'date')
            print(f"  📅 日期: {paragraph.text[:30]}")
        
        elif para_type == 'attachment_marker':
            # 附件标记：左对齐顶格、3号黑体
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            paragraph.paragraph_format.first_line_indent = Pt(0)
            paragraph.paragraph_format.left_indent = Pt(0)
            
            for run in paragraph.runs:
                run.font.name = '黑体'
                if run._element.rPr is not None:
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
                run.font.size = Pt(16)
                run.font.bold = True
                run.font.italic = False
                run.font.color.rgb = RGBColor(0, 0, 0)
            
            print(f"  📎 附件标记: {paragraph.text[:30]}")
        
        else:  # body
            apply_paragraph_format(paragraph, 'body')
    
    # 打印统计信息
    print(f"\n  📊 格式化统计:")
    for ptype, count in sorted(type_counts.items()):
        print(f"     {ptype}: {count} 个")


def llm_format_document(input_path):
    """LLM 增强格式化主函数"""
    try:
        print(f"\n🤖 [LLM模式] 正在处理: {os.path.basename(input_path)}")
        print("━" * 50)
        
        # 0. 检查 Ollama 连接
        print("  🔍 检查 Ollama 服务...")
        client = OllamaClient()
        success, message = client.check_connection()
        print(f"     {message}")
        
        if not success:
            raise Exception("Ollama 连接失败，请确保 Ollama 已启动并安装了 qwen2.5:7b 模型")
        
        # 1. 读取文档
        print("  ⏳ 读取文档...")
        doc = Document(input_path)
        
        # 2. 提取纯文本（只提取非空段落）
        print("  📝 提取文档文本...")
        paragraphs_text = []
        
        for para in doc.paragraphs:
            # 跳过表格中的段落
            if has_table(para):
                continue
            
            # 跳过图片段落
            if has_image(para):
                continue
            
            text = para.text.strip()
            if text:
                paragraphs_text.append({
                    "index": len(paragraphs_text),
                    "content": text
                })
        
        if len(paragraphs_text) == 0:
            raise Exception("文档中没有有效文本内容")
        
        print(f"     提取了 {len(paragraphs_text)} 个有效段落")
        
        # 构建发送给 LLM 的文本
        document_text = "\n".join([f"{p['index']}: {p['content']}" for p in paragraphs_text])
        
        # 3. 调用 LLM 识别
        print("  🤖 调用本地 Qwen 模型分析文档结构...")
        print("     (这可能需要10-60秒，请耐心等待)")
        
        llm_result = client.analyze_document(document_text)
        
        print(f"  ✅ LLM识别完成")
        
        # 4. 验证 LLM 结果
        if not validate_llm_result(llm_result, doc):
            raise Exception("LLM识别结果验证失败")
        
        # 5. 设置页边距（GB/T 9704-2012标准）
        section = doc.sections[0]
        section.top_margin = Cm(3.7)
        section.bottom_margin = Cm(3.5)
        section.left_margin = Cm(2.8)
        section.right_margin = Cm(2.6)
        print("  ✅ 页边距: 上3.7cm 下3.5cm 左2.8cm 右2.6cm")
        
        # 6. 根据 LLM 结果应用格式
        print("  🎨 根据 LLM 识别结果应用格式...")
        apply_formats_by_llm(doc, llm_result)
        
        # 7. 保存文档
        dir_name = os.path.dirname(input_path)
        base_name = os.path.basename(input_path)
        output_path = os.path.join(dir_name, f"llm_{base_name}")
        
        print(f"  💾 保存文档...")
        doc.save(output_path)
        
        print("━" * 50)
        print(f"✅ [LLM模式] 处理完成！")
        print(f"📁 输出文件: {output_path}\n")
        
        return True
        
    except Exception as e:
        print(f"❌ [LLM模式] 处理失败: {str(e)}")
        import traceback
        traceback.print_exc()
        return False


def main():
    """命令行测试入口"""
    print("\n" + "=" * 50)
    print("  🤖 公文格式调整工具 - LLM增强版")
    print("=" * 50)
    print("\n使用方法：")
    print("  1. 拖拽Word文档到此窗口")
    print("  2. 按回车键开始处理")
    print("  3. 输入 'q' 退出程序")
    print("\n" + "=" * 50 + "\n")
    
    while True:
        try:
            user_input = input("📎 请拖拽Word文档到此处（或输入q退出）: ").strip()
            
            if user_input.lower() == 'q':
                print("\n👋 再见！\n")
                break
            
            # 处理路径
            file_path = user_input.strip('"').strip("'").strip()
            file_path = file_path.replace('\\ ', ' ')
            file_path = file_path.replace('\\(', '(').replace('\\)', ')')
            file_path = file_path.replace('\\[', '[').replace('\\]', ']')
            file_path = file_path.replace('\\&', '&')
            
            if not file_path:
                continue
            
            if not os.path.exists(file_path):
                print(f"❌ 文件不存在: {file_path}\n")
                continue
            
            if not file_path.lower().endswith('.docx'):
                print("❌ 只支持.docx格式的文件\n")
                continue
            
            # 处理文档
            success = llm_format_document(file_path)
            
            if success:
                print("━" * 50)
                print("✨ 可以继续处理下一个文档")
                print("━" * 50 + "\n")
        
        except KeyboardInterrupt:
            print("\n\n👋 程序已终止\n")
            break
        except Exception as e:
            print(f"\n❌ 发生错误: {str(e)}\n")


if __name__ == '__main__':
    main()
